from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = (
    PROJECT_ROOT
    / "docs"
    / "experiment_plans"
    / "DM_Filament_Experimental_Research_Plan.docx"
)

# narrative_proposal preset with one named localization override:
# Calibri -> Malgun Gothic for Korean glyph coverage. All geometry, sizes,
# spacing, colors, and table tokens follow the preset.
FONT_BODY = "Malgun Gothic"
FONT_MATH = "Cambria Math"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_FILL = "F4F6F9"
BLUE_FILL = "E8EEF5"
GOLD_FILL = "FFF6D8"
GOLD = "7A5A00"
WHITE = "FFFFFF"
GRID = "C8D0DA"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(
    run,
    *,
    name: str = FONT_BODY,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color: str, size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_border_left(paragraph, color: str, size: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = styles["Title"]
    title.font.name = FONT_BODY
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_BODY
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(18)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT_BODY
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def add_custom_numbering(doc: Document) -> tuple[int, int, int]:
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_start = max(existing_abstract, default=0) + 1
    num_start = max(existing_num, default=0) + 1

    def make_abstract(abstract_id: int, fmt: str, text: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        abstract.append(level)
        return abstract

    def append_instance(abstract_id: int, num_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    # OOXML requires every abstractNum before the first concrete num element.
    first_num_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(
        first_num_index,
        make_abstract(abstract_start, "bullet", "•"),
    )
    numbering.insert(
        first_num_index + 1,
        make_abstract(abstract_start + 1, "decimal", "%1."),
    )
    append_instance(abstract_start, num_start)
    append_instance(abstract_start + 1, num_start + 1)
    return num_start, num_start + 1, abstract_start + 1


def add_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing_num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc: Document, text: str, _num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_paragraph_border_left(p, BLUE, size=8)
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_bullet(doc: Document, text: str, _bullet_num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_paragraph_border_left(p, GRID, size=6)
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, bold=True, color=DARK_BLUE)
    text_run = p.add_run(text)
    set_run_font(text_run)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_formula(doc: Document, formula: str, explanation: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2 if explanation else 0)
    run = p.add_run(formula)
    set_run_font(run, name=FONT_MATH, size=12, color=INK)
    if explanation:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(explanation)
        set_run_font(r2, size=9.5, color=MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    header_fill: str = LIGHT_FILL,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9.5, bold=True, color=INK)

    for row_data in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=9.5)

    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label_run = paragraph.add_run("Page ")
    set_run_font(label_run, size=9, color=MUTED)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run = paragraph.add_run()
    set_run_font(begin_run, size=9, color=MUTED)
    begin_run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run = paragraph.add_run()
    set_run_font(instruction_run, size=9, color=MUTED)
    instruction_run._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run = paragraph.add_run()
    set_run_font(separate_run, size=9, color=MUTED)
    separate_run._r.append(separate)

    result_run = paragraph.add_run("1")
    set_run_font(result_run, size=9, color=MUTED)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run = paragraph.add_run()
    set_run_font(end_run, size=9, color=MUTED)
    end_run._r.append(end)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(3)
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.add_run("DM FILAMENT | PURGE & PROPERTY RESOLUTION")
    set_run_font(header_run, size=8.5, bold=True, color=MUTED)
    set_paragraph_border_bottom(header_p, GRID, size=4)

    # Named compatibility override: keep the footer empty. Word 16 can enter
    # an excessive repagination loop when this table-heavy document has footer
    # fields or text. The quiet running header still provides page furniture.


def build_document() -> Document:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    # Word 16 compatibility override: this table-heavy Korean document is
    # rendered with indented item/step blocks instead of list numbering.
    bullet_num_id = 0
    decimal_a_num_id = 0
    decimal_b_num_id = 0
    decimal_flow_num_id = 0

    # Cover: editorial_cover pattern.
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(14)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("EXPERIMENTAL RESEARCH PLAN")
    set_run_font(run, size=10.5, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("DM 필라멘트 Purge 보정 및\n최소 Property 해상도 실험 계획")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Dead-zone 정량화, Checkerboard 최소 표현 단위,\n"
        "3D Gradient Direction 실증"
    )

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(8)
    line.paragraph_format.space_after = Pt(34)
    set_paragraph_border_bottom(line, BLUE, size=10)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    set_run_font(meta.add_run("Project: DM-filament inverse design"), size=11, bold=True, color=INK)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(4)
    set_run_font(meta2.add_run("Target printer: Original Prusa XL"), size=10.5, color=MUTED)
    meta3 = doc.add_paragraph()
    meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta3.add_run("작성일: 2026-07-26"), size=10.5, color=MUTED)

    doc.add_page_break()

    doc.add_heading("문서 목적", level=1)
    add_body(
        doc,
        "본 문서는 DM 필라멘트를 이용한 3D 출력에서 Property 전환 시 발생하는 "
        "유효 dead zone을 정량화하고, 보정된 purge volume을 적용했을 때 구현 가능한 "
        "최소 XY Property 표현 단위와 3차원 gradient direction을 실험적으로 검증하기 "
        "위한 연구계획서다."
    )
    add_callout(
        doc,
        "핵심 연구 논리",
        "Dead-zone 보정 실험으로 전환별 최소 purge volume을 얻고, 이를 checkerboard에 "
        "적용해 최소 표현 단위를 결정한 뒤, 동일 단위를 3D gradient 구조물에 적용한다.",
        fill=BLUE_FILL,
    )

    doc.add_heading("전체 실험 구조", level=1)
    add_table(
        doc,
        ["단계", "핵심 질문", "주요 산출물"],
        [
            ["Phase A", "Property A에서 B로 안정화되기까지 필요한 압출량은 얼마인가?", "전환별 purge volume matrix"],
            ["Phase B", "보정 후 독립적으로 표현 가능한 최소 XY cell 크기는 얼마인가?", "방향별 minimum property unit"],
            ["Phase C", "최소 단위를 이용해 목표 3D gradient 방향을 재현할 수 있는가?", "3D property map 및 방향 오차"],
        ],
        [1300, 4300, 3760],
        header_fill=BLUE_FILL,
    )

    doc.add_heading("연구 가설", level=1)
    for item in (
        "H1. Property 전환 응답은 누적 압출 부피의 함수로 정량화할 수 있다.",
        "H2. 전환별 purge volume을 적용하면 미보정 조건보다 checkerboard cell의 순도와 경계 정확도가 개선된다.",
        "H3. 최소 표현 단위는 출력 방향과 toolpath scheduling에 따라 달라질 수 있다.",
        "H4. 최소 표현 단위로 이산화한 3D Property field는 목표 gradient 방향과 통계적으로 일치한다.",
    ):
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("공통 용어 및 판정 개념", level=1)
    add_table(
        doc,
        ["용어", "정의"],
        [
            ["Delay volume", "Property 변경 이후 출력에서 새로운 Property가 처음 검출될 때까지의 압출 부피"],
            ["Transition volume", "새 Property가 검출되기 시작한 시점부터 정상상태에 도달할 때까지의 부피"],
            ["Effective dead zone", "DM 필라멘트 경계 번짐, 공급 지연, hotend 체류, 용융 혼합 및 압력 안정화를 포함한 시스템 응답"],
            ["Required purge volume", "출력이 목표 Property의 허용 오차에 진입하고 안정적으로 유지되기까지 필요한 부피"],
            ["Minimum property unit", "사전 정의한 순도·대비·경계 기준을 만족하는 가장 작은 XY cell"],
        ],
        [2500, 6860],
        header_fill=LIGHT_FILL,
    )

    doc.add_page_break()
    doc.add_heading("Phase A. Dead-zone 및 Purge Volume 보정", level=1)

    doc.add_heading("A-1. 목적", level=2)
    add_body(
        doc,
        "DM 필라멘트의 Property가 A에서 B로 바뀐 입력 경계 이후, 출력 조성이 목표 B에 "
        "안정적으로 도달하기까지 소비되는 필라멘트 길이와 부피를 측정한다. 이 값은 단순한 "
        "hotend 내부 체적이 아니라 최종 공정 전체를 포함하는 유효 dead-zone 응답이다."
    )

    doc.add_heading("A-2. 시험 재료 선정", level=2)
    for item in (
        "1차 광학 calibration: Black -> White 및 White -> Black",
        "2차 조성 calibration: Property 차이가 가장 큰 두 조합",
        "3차 실사용 calibration: 실제 최적화 결과에서 빈번하게 나타나는 전환",
        "모든 방향을 개별 항목으로 취급하며 A->B와 B->A를 동일하다고 가정하지 않는다.",
    ):
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("A-3. Calibration DM 필라멘트", level=2)
    add_callout(
        doc,
        "필라멘트 시퀀스",
        "충분히 긴 Property A 정상구간 -> 가능한 한 선명한 A/B 경계 -> 충분히 긴 "
        "Property B 정상구간. 양방향 실험 시 B 정상구간 뒤에 B/A 경계와 A 정상구간을 "
        "추가하되, 두 전환 사이에 완전한 정상상태가 존재해야 한다.",
    )
    add_body(
        doc,
        "최종 workflow의 보정값이 목적이라면 입력 DM 필라멘트의 경계 번짐을 포함한 "
        "end-to-end 응답을 사용한다. 노즐 단독의 dead zone을 별도로 보고할 경우에는 출력 전 "
        "DM 필라멘트의 축방향 Property profile을 먼저 측정하고 입력 응답과 출력 응답을 분리한다."
    )

    doc.add_heading("A-4. 시편 및 G-code 설계", level=2)
    add_body(
        doc,
        "단일 레이어의 연속 serpentine strip을 사용한다. 각 측정 구간은 동일한 E 길이에 "
        "대응하도록 번호를 부여하며, 전환 구간에서 정지·retraction·유량 변경이 발생하지 않도록 한다."
    )
    add_callout(
        doc,
        "권장 pilot bin",
        "한 측정 bin을 E=2 mm로 설정하면 1.75 mm 필라멘트 기준 약 4.81 mm3에 대응한다. "
        "전환곡선이 너무 급하면 1 mm E, 너무 완만하면 5 mm E로 조정한다.",
        fill=GOLD_FILL,
    )
    add_formula(
        doc,
        "V_E = (pi * d_f^2 / 4) * deltaE",
        "V_E: 공급 필라멘트 부피, d_f: 실제 필라멘트 직경, deltaE: 누적 E 길이",
    )
    add_body(
        doc,
        "1.75 mm 필라멘트에서 V_E는 약 2.405 x deltaE mm3이다. 정밀 분석에는 실제 직경 "
        "평균을 적용하고, 가능하면 출력 전후 질량과 재료 밀도로 실제 토출 부피를 교차 확인한다."
    )

    doc.add_heading("A-5. 고정 조건", level=2)
    add_table(
        doc,
        ["변수", "기준 기록값", "관리 방법"],
        [
            ["Tool / nozzle", "T번호, 노즐 직경, 노즐 사용 이력", "시험 전 사진 및 설정 기록"],
            ["Thermal", "노즐·베드 온도", "출력 전 열적 정상상태 확보"],
            ["Flow", "체적 유량, 선속도, extrusion multiplier", "전환 동안 일정하게 유지"],
            ["Geometry", "layer height, line width, path orientation", "동일 G-code 사용"],
            ["Filament", "직경, 건조 조건, 제조 batch", "동일 batch 및 건조 protocol"],
            ["Motion", "retraction, acceleration, corner speed", "측정 구간에서 변경 금지"],
        ],
        [1900, 3400, 4060],
    )

    doc.add_heading("A-6. 실행 절차", level=2)
    steps_a = (
        "필라멘트를 건조하고 실제 직경을 여러 위치에서 측정해 평균과 표준편차를 기록한다.",
        "Property A를 충분히 압출해 정상상태 기준 색상 또는 Property 값을 확보한다.",
        "연속 serpentine calibration G-code를 시작하고 A/B 입력 경계를 기준 E=0으로 표시한다.",
        "정지 없이 Property B 정상상태까지 출력하며 각 공간 bin에 누적 E 범위를 대응시킨다.",
        "시편을 동일한 조명·카메라·노출 조건에서 촬영하거나 해당 Property 측정 장비로 분석한다.",
        "A->B 및 B->A를 각각 최소 5회 반복한다.",
        "기준 조건 완료 후 온도와 체적 유량의 저·중·고 조건에서 민감도 실험을 수행한다.",
    )
    for step in steps_a:
        add_numbered_paragraph(doc, step, decimal_a_num_id)

    doc.add_heading("A-7. 측정 및 정규화", level=2)
    add_body(
        doc,
        "색상 기반 분석에서는 고정 조명, 고정 white balance, 고정 노출 및 색상 기준표를 "
        "사용한다. 각 bin 중앙의 CIE L*a*b* 값을 얻고 A-B 색상 벡터에 투영하여 B 비율을 "
        "0~1로 정규화한다."
    )
    add_formula(
        doc,
        "C_B = ((Lab_sample - Lab_A) dot (Lab_B - Lab_A)) / ||Lab_B - Lab_A||^2",
        "C_B=0: Property A, C_B=1: Property B. 결과는 0~1 범위로 제한한다.",
    )
    add_body(
        doc,
        "색상이 실제 기능성 Property를 충분히 대표하지 않는 경우에는 Raman, FTIR, XRF, "
        "현미경, 전기저항, DMA 또는 nanoindentation 등 Property에 맞는 직접 측정을 병행한다."
    )

    doc.add_heading("A-8. 산출 지표와 Purge 판정", level=2)
    add_table(
        doc,
        ["지표", "의미", "사용"],
        [
            ["V05", "새 Property가 약 5% 검출되는 누적 부피", "Delay zone 종료"],
            ["V50", "전환 응답의 중심", "방향별 지연 비교"],
            ["V95", "목표 Property의 약 95% 도달 부피", "초기 purge 후보"],
            ["V99", "엄격한 정상상태 도달 부피", "고정밀 조건 후보"],
            ["V95 - V05", "혼합 transition 폭", "전환 선명도 비교"],
            ["V_purge", "허용 오차 진입 후 연속 bin에서 유지되는 최초 부피", "최종 G-code 설정"],
        ],
        [1400, 3920, 4040],
    )
    add_callout(
        doc,
        "권장 판정 기준",
        "색상은 목표 대비 deltaE00 <= 2 또는 연구 목적에 맞춘 임계값, 기능성 Property는 "
        "목표 대비 오차 <= 5%를 시작점으로 사용한다. 최소 3개 연속 bin에서 기준이 유지되어야 "
        "정상상태로 판정한다. 최종 V_purge는 반복값의 평균보다 상한 신뢰값을 사용하는 편이 안전하다.",
        fill=GOLD_FILL,
    )

    doc.add_heading("A-9. 최소 실행 매트릭스", level=2)
    add_table(
        doc,
        ["전환", "온도", "유량", "반복", "목적"],
        [
            ["Black -> White", "기준", "기준", "n>=5", "가혹한 광학 오염 조건"],
            ["White -> Black", "기준", "기준", "n>=5", "방향 비대칭 확인"],
            ["대표 A -> B", "기준", "기준", "n>=5", "실제 Property 조건"],
            ["대표 B -> A", "기준", "기준", "n>=5", "방향 비대칭 확인"],
            ["Worst direction", "저/기준/고", "기준", "각 n>=3", "온도 민감도"],
            ["Worst direction", "기준", "저/기준/고", "각 n>=3", "유량 민감도"],
        ],
        [1900, 1500, 1500, 1400, 3060],
    )

    doc.add_page_break()
    doc.add_heading("Phase B. Checkerboard 최소 표현 단위", level=1)

    doc.add_heading("B-1. 목적", level=2)
    add_body(
        doc,
        "노즐·선폭·레이어 높이의 고정된 하드웨어 한계를 전제로, purge 보정 후 독립적인 "
        "Property region으로 인정할 수 있는 최소 XY cell 크기를 결정한다."
    )

    doc.add_heading("B-2. Cell 크기 설계", level=2)
    add_body(
        doc,
        "nominal 선폭이 아니라 실제 측정된 평균 선폭 w를 기준으로 cell 크기를 설정한다. "
        "각 크기에서 최소 8x8 또는 10x10 checkerboard를 사용한다."
    )
    add_table(
        doc,
        ["배율", "Cell 크기", "w=0.45 mm 예시"],
        [
            ["8w", "8 x 실제 선폭", "3.60 mm"],
            ["6w", "6 x 실제 선폭", "2.70 mm"],
            ["4w", "4 x 실제 선폭", "1.80 mm"],
            ["3w", "3 x 실제 선폭", "1.35 mm"],
            ["2w", "2 x 실제 선폭", "0.90 mm"],
            ["1w", "1 x 실제 선폭", "0.45 mm"],
        ],
        [1800, 4100, 3460],
        header_fill=BLUE_FILL,
    )

    doc.add_heading("B-3. 실험군", level=2)
    add_table(
        doc,
        ["조건", "출력 전략", "검증 목적"],
        [
            ["Control", "Property 변경 후 purge 없이 다음 cell 출력", "dead-zone 영향의 baseline"],
            ["Calibrated purge", "모든 실제 Property 전환에 V_purge 적용", "전환 보정 효과"],
            ["Purge + grouped order", "같은 Property cell을 한 레이어에서 그룹화", "최적 scheduling의 실제 최소 단위"],
        ],
        [2100, 3940, 3320],
        header_fill=LIGHT_FILL,
    )
    add_callout(
        doc,
        "중요한 구분",
        "Raster 순서는 cell마다 전환하는 worst-case stress test다. Property-grouped 순서는 "
        "동일 geometry를 유지하면서 전환 횟수를 줄이는 실사용 최적 조건이다. 논문의 최소 표현 "
        "단위는 grouped 결과로 보고하고 raster 결과는 전환 내구성 비교군으로 제시한다.",
        fill=BLUE_FILL,
    )

    doc.add_heading("B-4. 방향성과 적층", level=2)
    for item in (
        "X 방향 toolpath checkerboard",
        "Y 방향 toolpath checkerboard",
        "가능하면 45도 방향 checkerboard",
        "광학 최소 단위는 neutral base 2~3 layers 위의 노출된 pattern layer 1개로 평가",
        "실제 기계적 Property 시편은 동일 패턴을 여러 층 적층하되 별도 결과로 구분",
    ):
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("B-5. 실행 절차", level=2)
    steps_b = (
        "기준 조건으로 single-line coupon을 출력하고 실제 평균 선폭 w를 측정한다.",
        "8w부터 1w까지 동일한 cell 수를 가진 checkerboard panel을 자동 생성한다.",
        "Control, calibrated purge, purge + grouped order 조건을 각각 출력한다.",
        "각 조건을 최소 3회 반복하고 X, Y 및 선택적으로 45도 방향을 분리한다.",
        "fiducial을 기준으로 이미지를 원근 보정하고 cell별 목표 mask를 생성한다.",
        "cell 중앙 순도와 전체 cell 면적 정확도를 각각 계산한다.",
        "사전 정의한 모든 판정 기준을 통과하는 가장 작은 cell을 minimum property unit으로 결정한다.",
    )
    for step in steps_b:
        add_numbered_paragraph(doc, step, decimal_b_num_id)

    doc.add_heading("B-6. 평가 지표", level=2)
    add_table(
        doc,
        ["평가 항목", "권장 초기 기준", "비고"],
        [
            ["Cell 중심 Property 오차", "<= 5%", "경계 영향을 제외한 내부 순도"],
            ["정확 분류 면적", ">= 90%", "전체 cell mask 기준"],
            ["A/B contrast 보존율", ">= 80%", "대형 cell contrast 대비"],
            ["경계 위치 오차", "<= 0.5w", "X/Y 방향별 계산"],
            ["반복성", "n>=3 모두 통과", "평균만으로 통과시키지 않음"],
        ],
        [2750, 2300, 4310],
    )
    add_formula(
        doc,
        "Minimum property unit = smallest cell satisfying every preregistered criterion",
        "최종 값은 X, Y 및 45도 방향별로 별도 보고한다.",
    )

    doc.add_page_break()
    doc.add_heading("Phase C. 최소 단위 기반 3D Gradient Direction 실증", level=1)

    doc.add_heading("C-1. 목적", level=2)
    add_body(
        doc,
        "Phase B에서 얻은 minimum property unit을 3차원 구조물의 논리적 cell로 사용하여, "
        "목표 Property field의 크기와 방향이 실제 출력물 내부에서 재현되는지를 검증한다."
    )

    doc.add_heading("C-2. 단계별 시편", level=2)
    add_table(
        doc,
        ["시편", "Target field", "검증 목적"],
        [
            ["Gx", "P(x): X 방향 선형 증가", "XY 내 한 축의 gradient 재현"],
            ["Gy", "P(y): Y 방향 선형 증가", "toolpath 방향성 비교"],
            ["Gz", "P(z): Z 방향 layer별 증가", "레이어 기반 gradient 재현"],
            ["Gxyz", "P(x,y,z): 공간 대각선 방향 증가", "최종 3D gradient direction 실증"],
        ],
        [1300, 3900, 4160],
        header_fill=BLUE_FILL,
    )
    add_formula(
        doc,
        "P(x,y,z) = P_min + (P_max - P_min) * (x + y + z) / (L_x + L_y + L_z)",
        "Gxyz 예시: 목표 gradient vector는 (1,1,1) 방향이다.",
    )

    doc.add_heading("C-3. Property 이산화", level=2)
    add_body(
        doc,
        "연속 gradient는 가용한 Property library를 이용해 cell 단위의 이산 단계로 변환한다. "
        "pilot은 5단계 또는 9단계로 수행하고, 성공 후 18색 또는 전체 Property 조합으로 확장한다."
    )
    add_callout(
        doc,
        "논리적 voxel",
        "XY 크기는 Phase B의 minimum property unit, Z 크기는 한 layer 또는 여러 layer 묶음으로 "
        "정의한다. 이 voxel은 Property 할당 단위이며 실제 toolpath는 노즐 선폭과 layer height를 따른다.",
    )

    doc.add_heading("C-4. Purge와 DM 필라멘트 동기화", level=2)
    add_body(
        doc,
        "Property가 실제로 바뀌는 event 앞에서만 전환별 purge를 수행한다. 동일 Property의 "
        "인접 region에는 purge를 삽입하지 않는다. 한 레이어에서 가능한 region은 Property별로 "
        "그룹화하여 전환 횟수를 최소화한다."
    )
    add_formula(
        doc,
        "E_DM,i = E_object,i + E_purge,i",
        "각 Property 구간의 DM 필라멘트 길이는 구조물 소비량과 그 Property 안정화를 위한 purge 소비량의 합이다.",
    )
    add_callout(
        doc,
        "동기화 실패 조건",
        "최종 G-code에 purge만 추가하고 DM 필라멘트 step length와 spiral mapping을 갱신하지 않으면 "
        "첫 purge 이후 모든 Property 위치가 구조물 region과 어긋난다.",
        fill=GOLD_FILL,
    )

    doc.add_heading("C-5. 3D 시편 측정", level=2)
    for item in (
        "XY, XZ, YZ 단면 및 목표 gradient 방향 단면을 절단·연마한다.",
        "각 단면의 Property map을 동일 좌표계로 등록한다.",
        "목표 field와 측정 field의 RMSE, R2, gradient 크기 오차를 계산한다.",
        "목표 gradient vector와 측정 gradient vector 사이의 각도 오차를 계산한다.",
        "전체 purge 폐기량, 전환 횟수 및 추가 출력 시간을 함께 보고한다.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_formula(
        doc,
        "theta = arccos((grad P_target dot grad P_measured) / (|grad P_target| |grad P_measured|))",
        "theta가 작을수록 의도한 3D gradient direction을 정확히 재현한 것이다.",
    )

    doc.add_heading("C-6. 권장 성공 기준", level=2)
    add_table(
        doc,
        ["항목", "Pilot 기준", "최종 기준 설정"],
        [
            ["Property RMSE", "전체 범위의 <= 10%", "측정 장비 오차를 반영해 사전 등록"],
            ["Gradient direction error", "<= 10 degrees", "응용 목적에 맞게 강화"],
            ["Cell target attainment", ">= 90%", "minimum unit 내부 기준과 통일"],
            ["시편 반복성", "n>=3", "방향별 독립 평가"],
        ],
        [2700, 2400, 4260],
    )

    doc.add_page_break()
    doc.add_heading("Purge Section 자동화 및 소프트웨어 요구사항", level=1)

    doc.add_heading("자동 처리 흐름", level=2)
    flow_steps = (
        "Layer x Region execution plan에서 시간순 Property event를 읽는다.",
        "현재 Property와 다음 Property를 비교해 실제 전환 여부를 판정한다.",
        "calibration matrix에서 V_purge(current->next)를 조회한다.",
        "출력물 bounds와 bed bounds를 이용해 purge tower 위치를 선택한다.",
        "현재 layer height와 line width로 필요한 serpentine 경로 길이를 계산한다.",
        "원본 E 모드와 위치를 보존하면서 tower 이동, purge, wipe 및 복귀 G-code를 삽입한다.",
        "동일 purge 소비량을 DM step length, optimization input 및 spiral mapping에 반영한다.",
        "3D preview에서 object, purge 및 feed 구간을 서로 다른 색으로 표시한다.",
    )
    for step in flow_steps:
        add_numbered_paragraph(doc, step, decimal_flow_num_id)

    doc.add_heading("Purge 경로 길이", level=2)
    add_formula(
        doc,
        "L_path = V_purge / (line width * layer height)",
        "직사각형 bead 근사. 실제 extrusion width 모델과 flow calibration으로 보정 가능하다.",
    )
    add_body(
        doc,
        "예를 들어 line width 0.45 mm, layer height 0.20 mm에서 30 mm3를 purge하려면 "
        "약 333.3 mm의 purge path가 필요하다. 이 경로는 tower 내부의 여러 serpentine lane으로 나눈다."
    )

    doc.add_heading("필수 안전 조건", level=2)
    for item in (
        "purge tower와 object의 XY 충돌 및 travel envelope 검사",
        "Prusa XL bed 범위와 tool별 접근 가능 영역 검사",
        "한 레이어에 필요한 purge 경로가 tower footprint를 초과할 경우 자동 확장 또는 오류 처리",
        "absolute/relative E와 G92 E 상태의 정확한 보존",
        "Z-hop, retraction, wipe 및 복귀 위치 검증",
        "전환 없는 레이어에서 tower 안정성을 유지할 sparse support 경로",
        "예상 purge 총량과 DM 필라멘트 총 길이의 사전 표시",
    ):
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("Calibration 데이터 구조", level=2)
    add_callout(
        doc,
        "권장 key",
        "tool, nozzle diameter, temperature, volumetric flow, from_property, to_property, "
        "purge_volume_mm3, purge_e_mm, replicate count, confidence bound, measurement method",
        fill=BLUE_FILL,
    )

    doc.add_heading("실험 품질관리 및 위험요인", level=1)
    add_table(
        doc,
        ["위험요인", "영향", "관리 방법"],
        [
            ["입력 DM 경계 자체의 번짐", "노즐 dead zone 과대평가", "end-to-end 값으로 명시하거나 입력 profile 별도 측정"],
            ["색상과 실제 Property의 불일치", "광학 결과의 물성 해석 오류", "대표 전환에 직접 Property 측정 병행"],
            ["유량/온도 변동", "전환곡선 이동", "로그 기록 및 안정구간 확보"],
            ["Checkerboard 출력 순서", "purge 횟수와 열이력 혼입", "raster와 grouped 조건 분리"],
            ["Purge tower 불안정", "충돌 및 출력 실패", "footprint·지지 경로·bed 위치 자동 검증"],
            ["E 동기화 누락", "전체 Property 위치 drift", "object+purge 소비량 통합 검증"],
        ],
        [2400, 3100, 3860],
        header_fill=GOLD_FILL,
    )

    doc.add_heading("최종 산출물", level=1)
    for item in (
        "Property 전환별 V05, V50, V95, V99 및 V_purge database",
        "온도·유량·전환 방향에 따른 sensitivity plot",
        "보정 전/후 checkerboard 비교 이미지와 방향별 minimum property unit",
        "3D Gx, Gy, Gz 및 Gxyz gradient 시편",
        "목표/측정 Property field 비교와 gradient direction error",
        "Purge tower 자동 생성 G-code 및 3D/spiral mapping preview",
        "총 purge volume, 출력 시간 및 DM 필라멘트 소비량 report",
    ):
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("결과 해석의 논리", level=1)
    add_callout(
        doc,
        "증명 구조",
        "Phase A는 Property 전환 보정 가능성을 증명한다. Phase B는 보정된 시스템이 구현할 수 있는 "
        "최소 XY Property 단위를 증명한다. Phase C는 그 단위를 이용해 목표 3D Property field와 "
        "gradient direction을 실제 구조물 내부에 구현할 수 있음을 증명한다.",
        fill=BLUE_FILL,
    )

    doc.add_heading("실험 기록 양식", level=1)
    add_table(
        doc,
        ["필드", "기록"],
        [
            ["Run ID", ""],
            ["Date / operator", ""],
            ["Tool / nozzle", ""],
            ["Filament batch / diameter", ""],
            ["From -> To Property", ""],
            ["Temperature / flow / speed", ""],
            ["Layer height / line width", ""],
            ["Measured V05 / V50 / V95 / V99", ""],
            ["Selected V_purge", ""],
            ["Image / raw data path", ""],
            ["Notes / anomaly", ""],
        ],
        [3300, 6060],
    )

    doc.add_heading("참고 자료", level=1)
    references = (
        "Prusa Research. Purging volumes (MMU). "
        "https://help.prusa3d.com/article/purging-volumes-mmu_125097",
        "Prusa Research. Wipe tower. "
        "https://help.prusa3d.com/article/wipe-tower_125010",
        "Prusa Research. Post-processing scripts. "
        "https://help.prusa3d.com/article/post-processing-scripts_283913",
        "Transition Behavior in Blended Material Large Format Additive Manufacturing. "
        "Polymers, 18(2), 178. https://www.mdpi.com/2073-4360/18/2/178",
    )
    for ref in references:
        add_bullet(doc, ref, bullet_num_id)

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    raw_path = OUTPUT_PATH.with_suffix(".raw.docx")
    document.save(raw_path)

    # A clean reopen/save normalizes OOXML ordering for Microsoft Word 16.
    # Strip every footer child except pPr to avoid its repagination loop.
    normalized = Document(raw_path)
    for section in normalized.sections:
        for paragraph in section.footer.paragraphs:
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
    normalized.save(OUTPUT_PATH)
    raw_path.unlink()
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
