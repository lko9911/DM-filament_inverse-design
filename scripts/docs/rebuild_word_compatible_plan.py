from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE = (
    ROOT
    / "out"
    / "docx_qa"
    / "dm_experiment_plan"
    / "diagnostic_no_numbering_footer.docx"
)
OUTPUT = (
    ROOT
    / "docs"
    / "experiment_plans"
    / "DM_Filament_Experiment_Plan_Compatible.docx"
)

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)


def set_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 25, DARK_BLUE, 0, 8),
        ("Subtitle", 13, GRAY, 0, 16),
        ("Heading 1", 16, BLUE, 15, 8),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 3),
    ):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = style_name.startswith("Heading")


def copy_paragraph(source: Paragraph, target: Document) -> None:
    has_page_break = bool(source._p.xpath('.//w:br[@w:type="page"]'))
    if has_page_break:
        target.add_page_break()
        return

    text = source.text.strip()
    if not text:
        return

    source_style = source.style.name if source.style is not None else "Normal"
    allowed_styles = {"Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"}
    style = source_style if source_style in allowed_styles else "Normal"
    paragraph = target.add_paragraph(style=style)

    if style in {"Title", "Subtitle"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif source.alignment is not None:
        paragraph.alignment = source.alignment

    run = paragraph.add_run(text)
    set_font(run)

    if source_style == "Caption":
        run.italic = True
        run.font.color.rgb = GRAY


def copy_table(source: Table, target: Document) -> None:
    rows = len(source.rows)
    cols = len(source.columns)
    table = target.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, source_row in enumerate(source.rows):
        for col_index, source_cell in enumerate(source_row.cells):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            text = "\n".join(
                paragraph.text.strip()
                for paragraph in source_cell.paragraphs
                if paragraph.text.strip()
            )
            run = paragraph.add_run(text)
            set_font(run, size=9.2, bold=row_index == 0)
            if row_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    target.add_paragraph().paragraph_format.space_after = Pt(1)


def main() -> None:
    source = Document(SOURCE)
    output = Document()
    configure_document(output)

    page_break_count = 0
    for block in iter_blocks(source):
        if isinstance(block, Paragraph):
            is_break = bool(block._p.xpath('.//w:br[@w:type="page"]'))
            # Skip the Phase A -> B break if it would create a nearly blank page.
            if is_break:
                page_break_count += 1
                if page_break_count == 3:
                    continue
            copy_paragraph(block, output)
        else:
            copy_table(block, output)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    output.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
